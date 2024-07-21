from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

# 创建一个新的Word文档
doc = Document()

# 设置页面边距
sections = doc.sections
for section in sections:
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

# 标题页
doc.add_paragraph("表 1: 敘述性統計分析表", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("（N=45）", style='Subtitle').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 添加表格
table = doc.add_table(rows=1, cols=3)

# 设置表头
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '主要分類項目'
hdr_cells[1].text = '人數'
hdr_cells[2].text = '%'

# 样本数据
data = [
    ['性別', '', ''],
    ['生理男', '26', '57.8%'],
    ['女', '19', '42.2%'],
    ['年齡', '', ''],
    ['20 歲(含)以下', '5', '11.1%'],
    ['21-30 歲', '22', '48.9%'],
    ['31-40 歲', '11', '24.4%'],
    ['41-50 歲', '4', '8.9%'],
    ['51歲以上', '3', '6.7%'],
    ['教育程度', '', ''],
    ['國中(含)以下', '0', '0%'],
    ['高中職', '1', '2.2%'],
    ['大專院校', '27', '60.0%'],
    ['研究所', '17', '37.8%'],
    ['博士生', '0', '0%'],
    ['月收入', '', ''],
    ['20000(含)以下','14','31.1%'],
    ['20001-30000', '5', '11.1%'],
    ['30001-40000', '6', '13.3%'],
    ['40001-50000', '10', '22.2%'],
    ['50001以上', '10', '22.2%'],
]

# 插入数据到表格
for row_data in data:
    row_cells = table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row_cells[i].text = cell_data
        row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if i > 0 else WD_PARAGRAPH_ALIGNMENT.LEFT

# 设置表格样式
table.style = 'Table Grid'
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            run = paragraph.runs
            if run:
                run[0].font.size = Pt(12)

# 保存文档
doc.save('C:/Users/Leon/Desktop/Demographic_Variables_APA7.docx')

